import google.generativeai as genai
from pathlib import Path
import time
import os # Added for potentially getting API key from environment
from docx import Document
import re
import fitz  # PyMuPDF

# --- Function Definition ---
def getScores(api_key: str, folder_path: str) -> None:
    """
    Goes through all PDF files in the specified folder, analyzes each using Gemini
    with a fixed prompt, and saves the responses to text files.

    Args:
        api_key: Your Google Gemini API key.
        folder_path: Path to the folder containing PDF files.
    """
    try:
        genai.configure(api_key=api_key)
        print("Gemini API configured.")
    except Exception as e:
        print(f"Error configuring Gemini API: {e}")
        return

    folder = Path(folder_path.strip())
    if not folder.is_dir():
        print(f"Error: Folder not found at '{folder_path}'")
        return

    # Create analysis_output folder
    output_folder = Path(folder) / "analysis_output_3"
    output_folder.mkdir(exist_ok=True)
    print(f"Created/accessed output folder: {output_folder}")

    pdf_files = list(folder.glob("*.pdf"))
    if not pdf_files:
        print("No PDF files found in the specified folder.")
        return

       # 7. Define the analysis prompt
    prompt = """ You are tasked with analyzing a city's Climate Action Plan (CAP) or related report. Based on the content, assess whether it addresses key areas across political commitment, governance, stakeholder engagement, emissions data, risk assessments, strategies, equity, and monitoring.
        Please respond using Yes/No and provide brief justifications or references where applicable. Where methods, tools, or stakeholder names are mentioned, list or summarize them clearly. Provide a comprehensive analysis along with scores. A "Yes" must be given a score of 1 and a "No" must be given a score of 0.
        🔹 1. First Steps
        Has the city explicitly secured high-level political commitment to climate action or sustainability, endorsing a standard such as the Paris Agreement and defining ambitious goals?
        Has the city explicitly established and publicly communicated a clear, long-term vision for its low-carbon, climate-resilient, or sustainable future?
        🔹 2. Governance
        Has the city explicitly established clear cross-departmental governance structures or legally binding frameworks to ensure accountability and coordination for CAP implementation or a similar effort?
        🔹 3. Stakeholder & Community Engagement
            3.1 Identifying Priority Stakeholders:
        Does the report identify groups most impacted by climate change or sustainability issues (e.g., children, women, disabled, marginalized, frontline communities)?
        Are groups mentioned above excluded from previous engagement processes acknowledged?
        3.2 Engagement & Collaboration:
        Does the report mention planned engagement with the private sector, national/regional governments, or other stakeholders?
        Has the city identified influential actors supportive of its climate or sustainability plans?
            3.3 Engagement Methods:
        Have key stakeholders been integrated through long-term engagement throughout planning and implementation? If yes, which methods and groups?
        Has the broader public been engaged via surveys, consultations, summits, etc.?
        🔹 4. GHG Emissions Inventory
        Is the city measuring GHG emissions? If yes, what methodology is used?
        Does it use a specific tool for inventory management and reporting?
        Are the emissions inventory calculations publicly published?
        🔹 5. Sustainability Risk Assessment
        A sustainability risk assessment seeks to understand the likelihood of current and future sustainability hazards and the potential impacts of these hazards on cities and their inhabitants.A special case is a Climate Change Risk Assessment (CCRA). A climate hazard assessment identifies probability, intensity, and timescale. A climate impact assessment identifies impacts on people, infrastructure, and services.
        Has the city conducted a climate or sustainability hazard assessment (probability, intensity, timescale)?
        Has the city conducted a climate or sustainability impact assessment (on people, infrastructure, services)?
        Has the city conducted a full CCRA or related assessment?
        Has the CCRA or related assessment been outsourced? If so, to whom?
        Has the CCRA or related assessment been updated or scheduled for renewal?
        Has the CCRA or related assessment included interdependent risks and adaptive capacity analysis?
        Has the CCRA or related assessment been made public?
        🔹 6. City Needs Assessment
        Has the city analyzed socioeconomic context, environmental quality, and alignment with SDGs (Sustainable Development Goals) through strategic appraisal?
        Has it assessed city-wide priorities that climate or sustainability actions could address?
        🔹 7. Strategy Identification Mitigation Strategies:
            7.1 Mitigation Strategies
        Has the city defined a planning horizon for its climate or related scenarios?
        A BAU scenario is an emissions projection assuming no additional actions are implemented beyond current policies. Has a Business-As-Usual (BAU) forecast been included?
        Are modeling tools used for scenario development?
        Are mitigation projections based on current plans available?
        Is there an ambitious, long-term mitigation scenario?
        7.2 Adaptation Strategies:
        Has the city identified the root causes of climate or sustainability risks?
        Reactive adaptation fights the immediate negative consequences of climate or sustainability-related hazards, protecting quality of life and the city's systems during climate or sustainability-related disasters and restoring them afterwards. Are any reactive adaptation plans addressed? 
        Preventative adaptation reduces the negative consequences of climate or sustainability-related hazards, aiming to protect quality of life and city systems to avoid those hazard events becoming disasters. Are any preventive adaptation plans included? 
        Transformative adaptation tackles the root causes of climate or sustainability risks, making climate or sustainability-related hazards less likely or severe through fundamental changes to the city's fabric and systems. Are any transformative adaptation plans included? 
        🔹 8. Action Prioritization & Detailing
        Has a longlist of potential actions been developed from an evidence base?
        Has a shortlist of high-priority actions been defined using specific criteria/tools (e.g., ASAP, AMIA, cost-benefit)?
        Does the plan assess the fit of actions within broader city agendas?
        Is there evidence of inclusive stakeholder engagement in prioritization?
        Has the city adopted a flexible, iterative planning process?
        🔹 9. Equity & Inclusivity
            9.1 Stakeholder Inclusion:
        Has the city included a diverse set of stakeholders in planning?
        9.2 Needs & Vulnerability Assessment
        Has the city identified vulnerable groups and reasons for vulnerability?
        Has a comprehensive needs assessment been done?
        9.3 Distributed Impact Analysis:
        Are equity impacts and challenges of actions analyzed?
        Has the city used needs/stakeholder findings to guide climate or sustainability-related actions?
        9.4 Monitoring Equity
        Is a Monitoring, Evaluation, and Reporting (MER) system used to track equity outcomes?
        🔹 10. Monitoring, Evaluation & Reporting (MER) 
        Indicators are measurable variables used to evaluate progress or success, defined at three levels: Outputs (direct products), Outcomes (immediate changes resulting from actions), and Impacts (long-term effects contributing toward broader goals).
        10.1 Integration with City Systems:
        Are existing climate or sustainability plans and tracking mechanisms referenced?
        Are inclusivity, public reporting, and data systems discussed?
        10.2 Governance & Stakeholders:
        Are key stakeholders identified in MER? If yes, which ones? 
        10.3 Defining Indicators:
        Are clear indicators set for each action (output, outcome, impact)?
        Are GHG, risk, and co-benefits included?
        10.4: Data Collection
        Has the city identified data sources, ownership, collection methods, and reporting responsibilities?
        10.5: Updating
        Does the city explicitly commit to periodic evaluations of actions from CAP or related efforts against clearly defined baselines and targets?

        The scoring can be done in the way shown below:
        Counting the Yes/No Questions:
        First Steps: 
        Political commitment? (1), Long-term vision? (1) 
        Section 1 Total: 2 points
        Governance: 
        Governance structures/frameworks? (1)
        Section 2 Total: 1 point
        Stakeholder & Community Engagement: 
        Impacted groups identified? (1), Excluded groups acknowledged? (1), Planned engagement mentioned? (1), Influential actors identified? (1), Key stakeholders integrated long-term? (1), Broader public engaged? (1)
        Section 3 Total: 6 points
        GHG Emissions Inventory: 
        Measuring GHG emissions? (1), Specific tool used? (1), Calculations publicly published? (1)
        Section 4 Total: 3 points
        Sustainability Risk Assessment: 
        Hazard assessment? (1), Impact assessment? (1), Full CCRA conducted? (1), Outsourced CCRA? (1), Updated or scheduled renewal? (1), Interdependent risks/adaptive capacity included? (1), Publicly available? (1)
        Section 5 Total: 7 points
        City Needs Assessment: 
        Socioeconomic context analyzed? (1), City-wide priorities assessed? (1)
        Section 6 Total: 2 points
        Strategy Identification: 
        Planning horizon defined? (1), BAU forecast included? (1), Modeling tools used? (1), Mitigation projections available? (1), Ambitious scenario defined? (1), Root causes identified? (1), Reactive adaptation included? (1), Preventive adaptation included? (1), Transformative adaptation included? (1)
        Section 7 Total: 9 points
        Action Prioritization & Detailing: 
        Longlist developed? (1), Shortlist defined? (1), Fit within agendas assessed? (1), Inclusive engagement in prioritization? (1), Flexible iterative planning adopted? (1)
        Section 8 Total: 5 points
        Equity & Inclusivity: 
        Diverse stakeholders included? (1), Vulnerable groups identified? (1), Comprehensive needs assessment? (1), Equity impacts analyzed? (1), Findings guide actions? (1), MER for equity used? (1)
        Section 9 Total: 6 points
        Monitoring, Evaluation & Reporting (MER): 
        Existing plans referenced? (1), Inclusivity/public reporting/data systems discussed? (1), Key stakeholders identified? (1), Clear indicators set? (1), GHG/risk/co-benefits included? (1), Data sources/methods identified? (1), Periodic evaluations committed? (1)
        Section 10 Total: 7 points
        Calculating the Total Maximum Score:
        Adding the maximum points from each section: 2 + 1 + 6 + 3 + 7 + 2 + 9 + 5 + 6 + 7 = 48 points
        Based on the provided structure and counting each distinct Yes/No question as one point, the maximum score any given report can get is 48.
        """

    model = genai.GenerativeModel(model_name="gemini-2.5-pro-preview-03-25")

    for pdf_path in pdf_files:
        print(f"\nProcessing '{pdf_path.name}'...")
        try:
            # Extract and clean text from PDF deterministically
            pdf_text = extract_text_from_pdf(str(pdf_path))
            clean_text = clean_extracted_text(pdf_text)
            print(f"Extracted and cleaned text from {pdf_path.name} (first 500 chars):\n{clean_text[:500]}\n---")

            print("Sending request to Gemini for analysis...")
            full_input = f"{prompt.strip()}\n\n---\n\n{clean_text.strip()}"
            response = model.generate_content(
                [full_input],
                generation_config={
                    "temperature": 0.0,
                    "top_p": 0.0
                }
            )

            if response.text:
                output_file = output_folder / f"{pdf_path.stem}_analysis.docx"
                doc = Document()
                doc.add_heading(f"Analysis for {pdf_path.name}", level=1)
                bold_pattern = re.compile(r"\*\*(.+?)\*\*")
                for line in response.text.strip().split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    p = doc.add_paragraph()
                    last_idx = 0
                    for match in bold_pattern.finditer(line):
                        if match.start() > last_idx:
                            p.add_run(line[last_idx:match.start()])
                        bold_run = p.add_run(match.group(1))
                        bold_run.bold = True
                        last_idx = match.end()
                    if last_idx < len(line):
                        p.add_run(line[last_idx:])
                doc.save(output_file)
                print(f"Analysis saved to: {output_file.name}")
            else:
                print("[No text content received in the response]")

        except Exception as e:
            print(f"\nAn error occurred during processing or analysis: {e}")

def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extracts clean, readable text from a PDF using PyMuPDF.
    Args:
        pdf_path: Path to the PDF file.
    Returns:
        Full extracted text as a string.
    """
    path = Path(pdf_path)
    if not path.is_file() or path.suffix.lower() != ".pdf":
        raise ValueError(f"Invalid PDF file: {pdf_path}")
    text_parts = []
    with fitz.open(pdf_path) as doc:
        for page in doc:
            page_text = page.get_text("text")  # Preserve layout where possible
            text_parts.append(page_text.strip())
    full_text = "\n\n".join(text_parts)
    return full_text.strip()

def clean_extracted_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)               # collapse spaces
    text = re.sub(r"\n{2,}", "\n", text)              # remove multiple line breaks
    text = re.sub(r" +\n", "\n", text)                # remove trailing spaces on lines
    text = text.encode("ascii", "ignore").decode()    # remove non-ASCII chars
    return text.strip()


# --- Example Usage (How to call the function) ---
if __name__ == "__main__":
    my_api_key = "AIzaSyDZ1hkeltOGVCVMT6h_lRZGNpyfIgwDOeY"
    folder_path_input = "/Users/satendergunwal/Desktop/CGithub-SGunwal/CLEAN/sustainable_maturity_mapping/single_test_report" #"community_reports/GRC_Reports" #"test_folder" #input("Please enter the full path to the folder containing PDF files: ")

    if my_api_key:
        print("\nStarting batch analysis process...")
        getScores(my_api_key, folder_path_input)
        print("\nBatch processing completed.")
    else:
        print("Cannot proceed without a Gemini API Key.")