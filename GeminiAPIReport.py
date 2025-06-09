import google.generativeai as genai
from pathlib import Path
import time
import os
from docx import Document
import re
import json
import pandas as pd
from json import JSONDecodeError
from ast import literal_eval

# --- Helper to clean and extract JSON ---
def extract_json_from_response(text: str):
    text = text.strip()

    # Try to extract first valid JSON array using regex
    json_array_match = re.search(r'\[\s*{.*?}\s*]', text, re.DOTALL)
    if json_array_match:
        try:
            return json.loads(json_array_match.group(0))
        except Exception:
            pass

    # Fallback: Remove markdown if any
    if text.startswith("```json"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    text = text.strip()

    try:
        return json.loads(text)
    except JSONDecodeError:
        try:
            return literal_eval(text)
        except Exception:
            return None

# --- Main Scoring Function ---
def getScores(api_key: str, folder_path: str) -> None:
    try:
        genai.configure(api_key=api_key)
        print("Gemini API configured.")
    except Exception as e:
        print(f"Error configuring Gemini API: {e}")
        return

    folder = Path(folder_path.strip())
    if not folder.is_dir():
        print(f"Error: Folder not found at '{folder_path}'")
        return

    output_folder = folder / "analysis_output"
    output_folder.mkdir(exist_ok=True)
    print(f"Created/accessed output folder: {output_folder}")

    pdf_files = list(folder.glob("*.pdf"))
    if not pdf_files:
        print("No PDF files found in the specified folder.")
        return

    summary_data = []

       # 7. Define the analysis prompt
    prompt = """
You are provided with a set of Yes/No questions below, which you must answer strictly using information explicitly stated in a Climate Action Plan (CAP) or similar sustainability document, referred to as the "Source Document."
Instructions for Answering Each Question (JSON Output Required):
For each question, strictly follow these two steps, formatting your responses in the provided JSON structure:
•	Step 1: Identify Relevant Text Snippet
Quote the exact text from the Source Document explicitly addressing the question. If the Source Document does not explicitly contain relevant information, write:
"None explicitly stated."
•	Step 2: Provide Your Answer and Justification
o	Answer "Yes" only if the quoted snippet explicitly provides the exact information requested.
o	Otherwise, answer "No".
o	Provide a brief justification explicitly referencing the quoted snippet and page number.

Important Guidelines:
•	Explicit evidence means information clearly stated in the Source Document that directly addresses the substance of the question. The wording need not exactly match the question, but the meaning must clearly and explicitly correspond.
•	Only information clearly and explicitly present in the Source Document should be used. If the Source Document does not clearly and explicitly provide information addressing the substance of the question, answer "No."
•	Avoid inferring, interpreting, or extrapolating information.
•	Remain objective, neutral, and consistent across all answers.
•	Assign scores strictly: "Yes" = 1 point, "No" = 0 points.

✅ JSON Output Example of a "Yes" Answer
{
  "question_id": "Example-Yes-1",
  "question_text": "Has the city explicitly adopted the GPC (BASIC or BASIC+) to structure its emissions inventory?",
  "relevant_snippet": "The source document explicitly states that it is consistent with the Global Protocol for Community-Scale Greenhouse Gas Inventories BASIC protocol.",
"page_no": "20",
  "answer": "Yes",
  "justification": "This is consistent with the Global Protocol for Community-Scale Greenhouse Gas Inventories BASIC protocol, and satisfies the Greenhouse Gas Emissions Inventory reporting requirement in the GCoM Common Reporting Framework.",
  "score": 1
}

❌ JSON Output Example of a "No" Answer
{
  "question_id": "Example-No-1",
  "question_text": "Has the city explicitly adopted the GPC (BASIC or BASIC+) to structure its emissions inventory?",
  "relevant_snippet": "N/A",
"page_no": "N/A",
  "answer": "No",
  "justification": "The document does not explicitly identify vulnerable populations. It only mentions general 'community engagement.'",
  "score": 0
}

Important Instruction: Please respond with a single valid JSON array only, no markdown, no commentary, no formatting. Do not wrap the output in backticks or explain anything.

1. Emissions Inventory (Max: 7 points)
1.1 Does the source document explicitly mention that the city has an emissions inventory?  (Yes=1 / No=0)
1.2 Does the source document explicitly mention if the city adopted the GPC (BASIC or BASIC+) to structure its emissions inventory? (Yes=1 / No=0)
1.3 Does the source document explicitly mention if the city's emissions inventory explicitly include Scope 1 stationary energy emissions (on-site fuel combustion)? (Yes=1 / No=0)
1.4 Does the source document explicitly mention if the city's emissions inventory explicitly include Scope 2 stationary energy emissions (grid-supplied electricity, heat, or steam)? (Yes=1 / No=0)
1.5 Does the source document explicitly mention if the city's emissions inventory explicitly state whether Scope 3 stationary energy emissions (e.g., upstream energy losses, imported fuels) are included or not? (Yes=1 / No=0)
1.6 Does the source document explicitly mention whether the emissions inventory disaggregates stationary energy emissions by sub-sector (e.g., residential, commercial, industrial)?  (Yes=1 / No=0)
1.7 Does the source document explicitly mention whether the emissions inventory disaggregates stationary energy emissions by fuel type (e.g., electricity, natural gas, coal, renewables)? (Yes=1 / No=0)


2. Strategy Identification (Max: 7 points)
2.1 Does the source document explicitly mention the use of emissions modeling or scenario planning for stationary energy emissions reduction strategies? (Yes=1 / No=0)
2.2	Does the source document explicitly identify technological solutions (e.g., renewable energy, retrofitting, heat pumps)? (Yes=1 / No=0)
2.3	Does the source document explicitly identify non-technological solutions (e.g., policy changes, behavior change programs)? (Yes=1 / No=0)
2.4	Does the source document explicitly set GHG reduction targets specifically for stationary energy sub-sectors (e.g., residential buildings, commercial buildings, industrial facilities, or energy utilities)? (Yes=1 / No=0)
2.5	Does the source document explicitly address residual stationary energy emissions (e.g., offsets, negative emission strategies)? (Yes=1 / No=0)
2.6	Does the source document explicitly assess how climate risks (e.g., heatwaves, flooding) impact the stationary energy sector (demand or reliability)? (Yes=1 / No=0)
2.7	Does the source document explicitly identify adaptation strategies (e.g., resilient infrastructure, grid reliability enhancements) addressing climate risks specific to the stationary energy sector? (Yes=1 / No=0)


3. Action Prioritization & Detailing (Max: 6 points)
3.1 Does the source document explicitly have actions to reduce stationary energy emissions? (Yes=1 / No=0)
3.2	Does the source document explicitly prioritize stationary energy actions based on quantified or qualitatively stated GHG emissions reduction potential? (Yes=1 / No=0)
3.3	Does the source document explicitly document feasibility assessments (technical, financial, institutional, or political) for prioritized stationary energy actions? (Yes=1 / No=0)
3.4	Does the source document explicitly address equity by considering vulnerable communities (e.g., low-income housing) when prioritizing stationary energy actions? (Yes=1 / No=0)
3.5	Does the source document explicitly define the expected GHG or climate adaptation impact for each prioritized stationary energy action? (Yes=1 / No=0)
3.6	Does the source document explicitly identify responsible departments or governance structures for implementing prioritized stationary energy actions? (Yes=1 / No=0)

4. Monitoring, Evaluation & Reporting (MER) (Max: 5 points)
4.1	Does the source document explicitly establish Key Performance Indicators (KPIs) to monitor stationary energy actions (e.g., energy savings, buildings retrofitted, emissions reduced)? (Yes=1 / No=0)
4.2	Does the source document explicitly describe regular evaluations (annual, biennial, mid-term) of stationary energy actions? (Yes=1 / No=0)
4.3	Does the source document explicitly commit to publicly reporting progress on stationary energy actions (e.g., annual reports, dashboards)? (Yes=1 / No=0)
4.4	Does the source document explicitly state that monitoring data is used to revise or update the CAP or stationary energy strategies? (Yes=1 / No=0)
4.5	Does the source document explicitly describe data quality assurance procedures (validation or verification of stationary energy emission data)? (Yes=1 / No=0)

The scoring can be done in the way shown below
Counting the Yes/No Questions:
Energy Emissions Inventory:
1.1 City have emissions inventory? (1)
1.2 GPC Adoption? (1)
1.3 Inclusion of Scopes 1? (1)
1.4 Inclusion of Scopes 2? (1)
1.5 Inclusion of Scopes 3? (1)
1.6 Sub-sector Disaggregation? (1) 
1.7 Fuel Type Disaggregation? (1)
= 7 points

Strategy Identification:
2.1 Use of Modeling/Scenario Planning? (1)
2.2 Identification of Technological Strategies? (1)
2.3 Identification of Non-technological strategies? (1)
2.4 Sector-Specific Targets? (1)
2.5 Residual Emissions and Offsets? (1)
2.6 Climate Risk Impact on Energy Sector? (1)
2.7 Adaptation Strategies (Reactive, Preventative, Transformative)?(1)
= 7 points

Action Prioritization & Detailing:
3.1 City have actions to reduce stationary energy emissions(1)
3.2 Prioritization Based on GHG Potential? (1)
3.3 Feasibility Assessment? (1)
3.4 Equity Consideration? (1)
3.5 Definition of GHG/Adaptation Impact? (1)
3.6 Responsible Departments Identified? (1)
= 6 points

Monitoring, Evaluation & Reporting (MER):
4.1 KPIs Established? (1)
4.2 Regular Evaluation Process? (1)
4.3 Public Progress Reporting? (1)
4.4 Use of Monitoring Data for Revisions? (1)
4.5 Data quality assurance procedures? (1)
= 5 point

Final Scoring:
•	Maximum score = 25
•	Total score = ___ / 25
"""

    model = genai.GenerativeModel(
            model_name="gemini-2.5-flash-preview-05-20",
            generation_config={
            "temperature": 0.0
            }
        )

    for pdf_path in pdf_files:
        print(f"\nProcessing '{pdf_path.name}'...")
        pdf_file = None

        try:
            pdf_file = genai.upload_file(path=pdf_path)
            print(f"Uploaded: {pdf_path.name}")

            while pdf_file.state.name == "PROCESSING":
                time.sleep(5)
                pdf_file = genai.get_file(name=pdf_file.name)

            if pdf_file.state.name != "ACTIVE":
                print(f"File not active. Skipping: {pdf_path.name}")
                genai.delete_file(name=pdf_file.name)
                continue

            print("Sending prompt to Gemini...")
            response = model.generate_content([prompt, pdf_file])

            json_output_file = output_folder / f"{pdf_path.stem}_analysis.json"
            raw_text = response.text.strip() if response.text else ""
            analysis_data = extract_json_from_response(raw_text)

            if analysis_data is None:
                print("❌ JSON parsing failed. Saving raw response.")
                with open(json_output_file, "w", encoding="utf-8") as f:
                    f.write(raw_text)
                continue
            else:
                with open(json_output_file, "w", encoding="utf-8") as f:
                    json.dump(analysis_data, f, indent=2)
                print(f"✅ Structured JSON saved: {json_output_file.name}")

            output_docx = output_folder / f"{pdf_path.stem}_analysis.docx"
            doc = Document()
            doc.add_heading(f"Analysis for {pdf_path.name}", level=1)

            category_scores = {
                "Emissions Inventory": 0,
                "Strategy Identification": 0,
                "Action Prioritization & Detailing": 0,
                "Monitoring, Evaluation & Reporting (MER)": 0
            }

            if isinstance(analysis_data, list):
                for item in analysis_data:
                    qid = item.get("question_id", "N/A")
                    qtext = item.get("question_text", "N/A")
                    snippet = item.get("relevant_snippet", "N/A")
                    page = item.get("page_no", "N/A")
                    ans = item.get("answer", "N/A")
                    justification = item.get("justification", "")
                    score = item.get("score", 0)

                    if qid.startswith("1."):
                        category_scores["Emissions Inventory"] += score
                    elif qid.startswith("2."):
                        category_scores["Strategy Identification"] += score
                    elif qid.startswith("3."):
                        category_scores["Action Prioritization & Detailing"] += score
                    elif qid.startswith("4."):
                        category_scores["Monitoring, Evaluation & Reporting (MER)"] += score

                    doc.add_paragraph(f"Question ID: {qid}")
                    doc.add_paragraph(f"Question: {qtext}")
                    doc.add_paragraph(f"Relevant Snippet: {snippet}")
                    doc.add_paragraph(f"Page Number: {page}")
                    doc.add_paragraph(f"Answer: {ans}")
                    doc.add_paragraph(f"Justification: {justification}")
                    doc.add_paragraph(f"Score: {score}")
                    doc.add_paragraph("-" * 40)

                doc.save(output_docx)
                print(f"📄 Word report saved: {output_docx.name}")

                summary_data.append({
                    "City": pdf_path.stem,
                    **category_scores,
                    "Total Score": sum(category_scores.values())
                })
            else:
                print("⚠️ Skipping Excel summary — analysis data not a list.")

        except Exception as e:
            print(f"\n🚨 Error during analysis: {e}")

        finally:
            if pdf_file and hasattr(pdf_file, 'name'):
                try:
                    genai.delete_file(name=pdf_file.name)
                    print(f"🧹 Deleted file from Gemini: {pdf_file.name}")
                except Exception as cleanup_error:
                    print(f"⚠️ Cleanup failed for {pdf_file.name}: {cleanup_error}")

    if summary_data:
        df = pd.DataFrame(summary_data)
        excel_path = output_folder / "analysis_scores.xlsx"
        df.to_excel(excel_path, index=False)
        print(f"\n✅ Excel summary saved at: {excel_path}")


# --- Entry Point ---
if __name__ == "__main__":
    my_api_key = "AIzaSyDZ1hkeltOGVCVMT6h_lRZGNpyfIgwDOeY"
    folder_path_input = "D:/All Municipality Reports/Final_Municipality_Reports_For_Analysis"

    if my_api_key:
        print("\n🚀 Starting batch analysis...")
        getScores(my_api_key, folder_path_input)
        print("\n✅ Batch processing completed.")
    else:
        print("❌ Gemini API key not provided.")
