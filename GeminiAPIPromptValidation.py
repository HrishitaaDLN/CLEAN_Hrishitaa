import google.generativeai as genai
from pathlib import Path
from docx import Document
import re
import time
import pandas as pd

# --- CUSTOM PROMPT --- (extracted for reusability)
GEMINI_PROMPT = """
You are an expert validator analyzing Climate Action Plans (CAPs) and C40 Framework requirements. Your task is to perform a comprehensive content validation between C40 guidance documents and assessment criteria.

OBJECTIVE:
Perform a detailed analysis of C40 Climate Action Planning Framework documents to:
1. Extract all key requirements, guidelines, and best practices
2. Create a comprehensive content validation matrix
3. Identify any gaps in the assessment questions

ANALYSIS STEPS:

1. CONTENT EXTRACTION
- Read the provided C40 PDF thoroughly
- Extract ALL key requirements, guidelines, criteria, and best practices
- Categorize them under relevant themes (e.g., Emissions Inventory, Strategy Development, Implementation, Monitoring)
- Note specific technical requirements, methodologies, or standards mentioned

2. SYSTEMATIC VALIDATION
For each extracted requirement:
- Map it against the 23 assessment questions
- Determine if it's fully, partially, or not covered
- Note specific question numbers that address each requirement
- Identify any missing aspects or gaps

3. OUTPUT GENERATION

Generate three detailed matrices:

Matrix 1: C40 Requirements Coverage
| C40 Requirement | Source Section | Requirement Type | Covered by Question(s)? | Question ID(s) | Coverage Level | Notes |
Example:
| GPC Protocol Usage | Emissions Inventory | Mandatory | Yes | 1.1 | Full | Explicitly covered in inventory question |

Matrix 2: Assessment Questions Validation
| Question ID | Question Text | C40 Requirements Covered | Missing Requirements | Suggested Improvements |
Example:
| 1.1 | GPC adoption question | GPC protocol requirement | None | Consider adding version specification |

Matrix 3: Gap Analysis
| Gap Area | Missing C40 Requirements | Importance Level | Suggested New Questions |
Example:
| Data Quality | Independent verification requirement | High | Add question about third-party verification |

ASSESSMENT QUESTIONS FOR REFERENCE:
1. Emissions Inventory 
1.1 Has the city explicitly adopted the GPC (BASIC or BASIC+) to structure its emissions inventory? 
1.2 Does the city's emissions inventory explicitly include Scope 1 stationary energy emissions (on-site fuel combustion)? 
1.3 Does the city's emissions inventory explicitly include Scope 2 stationary energy emissions (grid-supplied electricity, heat, or steam)? 
1.4 Does the city's emissions inventory explicitly state whether Scope 3 stationary energy emissions (e.g., upstream energy losses, imported fuels) are included or not? 
1.5 Does the emissions inventory explicitly disaggregate stationary energy emissions by sub-sector (e.g., residential, commercial, industrial)?
1.6 Does the emissions inventory explicitly disaggregate stationary energy emissions by fuel type (e.g., electricity, natural gas, coal, renewables)? 
2. Strategy Identification 
2.1 Does the city explicitly mention using emissions modeling or scenario planning for stationary energy emissions reduction strategies? 
2.2	Does the city explicitly identify technological solutions (e.g., renewable energy, retrofitting, heat pumps)? 

2.3	Does the city explicitly identify non-technological solutions (e.g., policy changes, behavior change programs)? 

2.4	Has the city explicitly set GHG reduction targets specifically for stationary energy sub-sectors (e.g., residential buildings, commercial buildings, industries, or energy utilities)? 

2.5	Has the city explicitly addressed residual stationary energy emissions (e.g., offsets, negative emission strategies)? 

2.6	Has the city explicitly assessed how climate risks (e.g., heatwaves, flooding) impact the stationary energy sector (demand or reliability)? 

2.7	Has the city explicitly identified adaptation strategies (e.g., resilient infrastructure, grid reliability enhancements) addressing climate risks specific to the stationary energy sector? 

3. Action Prioritization & Detailing 
3.1	Does the city explicitly prioritize stationary energy actions based on quantified or qualitatively stated GHG emissions reduction potential? 

3.2	Does the city explicitly document feasibility assessments (technical, financial, institutional, or political) for prioritized stationary energy actions? 

3.3	Does the city explicitly address equity by considering vulnerable communities (e.g., low-income housing) when prioritizing stationary energy actions?

3.4	For each prioritized stationary energy action, does the city explicitly define the expected GHG or climate adaptation impact? 

3.5	Does the city explicitly identify responsible departments or governance structures for implementing prioritized stationary energy actions?

4	Monitoring, Evaluation & Reporting (MER)

4.1	Does the city explicitly establish Key Performance Indicators (KPIs) to monitor stationary energy actions (e.g., energy savings, buildings retrofitted, emissions reduced)? 

4.2	Does the city explicitly describe regular evaluations (annual, biennial, mid-term) of stationary energy actions?

4.3	Does the city explicitly commit to publicly reporting progress on stationary energy actions (e.g., annual reports, dashboards)? 

4.4	Does the city explicitly state that monitoring data is used to revise or update the CAP or stationary energy strategies? 

4.5	Does the city explicitly describe data quality assurance procedures (validation or verification of stationary energy emission data)? 

SPECIFIC VALIDATION CRITERIA:
1. Completeness: Does the question capture all aspects of the C40 requirement?
2. Specificity: Is the question specific enough to assess compliance?
3. Measurability: Can the response be objectively evaluated?
4. Relevance: Does the question align with C40's intent?

IMPORTANT CONSIDERATIONS:
- Focus on both explicit and implicit requirements in C40 documents
- Consider technical, procedural, and governance aspects
- Note any regional or city-size specific requirements
- Identify both mandatory and recommended practices
- Consider alignment with international standards referenced by C40

Please provide detailed, actionable feedback that can be used to improve the assessment framework.
"""

# --- Function Definition ---
def getScores(api_key: str, folder_path: str) -> None:
    try:
        genai.configure(api_key=api_key)
        print(" Gemini API configured successfully.")
    except Exception as e:
        print(f" Error configuring Gemini API: {e}")
        return

    folder = Path(folder_path.strip())
    if not folder.is_dir():
        print(f" Folder not found at '{folder_path}'")
        return

    output_folder = folder / "analysis_output"
    output_folder.mkdir(exist_ok=True)
    print(f" Output folder ready: {output_folder}")

    pdf_files = list(folder.glob("*.pdf"))
    if not pdf_files:
        print(" No PDF files found.")
        return

    model = genai.GenerativeModel(model_name="gemini-2.5-pro-preview-03-25")

    for pdf_path in pdf_files:
        print(f"\n Processing: {pdf_path.name}")
        pdf_file = None

        try:
            pdf_file = genai.upload_file(path=pdf_path)
            print(f"Uploaded: {pdf_file.name}")

            # Wait for Gemini to process the file
            while pdf_file.state.name == "PROCESSING":
                print(" Waiting for Gemini to process file...")
                time.sleep(5)
                pdf_file = genai.get_file(name=pdf_file.name)

            if pdf_file.state.name != "ACTIVE":
                raise Exception(f"File processing failed. Final state: {pdf_file.state.name}")

            # Run analysis
            print(" Running content analysis...")
            response = model.generate_content([GEMINI_PROMPT, pdf_file])

            # Handle response
            if response.text:
                print(response.text)
                output_file = output_folder / f"{pdf_path.stem}_analysis.xlsx"
                save_text_as_excel(response.text.strip(), output_file)
                print(f" Analysis saved: {output_file.name}")
            else:
                print(" No text content received in the response.")

        except Exception as e:
            print(f" Error during analysis: {e}")

        finally:
            if pdf_file and hasattr(pdf_file, 'name'):
                try:
                    genai.delete_file(name=pdf_file.name)
                    print(f"Deleted file from Gemini: {pdf_file.name}")
                except Exception as cleanup_error:
                    print(f" Cleanup failed: {cleanup_error}")

# --- Helper to save Gemini output as Word document with bold formatting ---
def save_text_as_docx(text: str, output_path: Path) -> None:
    doc = Document()
    doc.add_heading("C40 Climate Action Plan Analysis", level=1)
    bold_pattern = re.compile(r"\*\*(.+?)\*\*")

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        last_idx = 0
        for match in bold_pattern.finditer(line):
            if match.start() > last_idx:
                p.add_run(line[last_idx:match.start()])
            p.add_run(match.group(1)).bold = True
            last_idx = match.end()
        if last_idx < len(line):
            p.add_run(line[last_idx:])
    doc.save(output_path)

# --- Helper to save Gemini output as Excel file ---
def save_text_as_excel(text: str, output_path: Path) -> None:
    # Initialize lists to store data for both tables
    coverage_matrix = []
    validation_table = []
    
    current_table = None
    
    # Parse the text and organize into dataframes
    lines = text.splitlines()
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if "Output Format 1:" in line or "C40 Key Parameter Coverage Matrix" in line:
            current_table = "coverage"
            continue
        elif "Output Format 2:" in line or "Enhanced Validation Table" in line:
            current_table = "validation"
            continue
            
        if current_table == "coverage" and ("\t" in line or "|" in line):
            # Handle both tab-separated and pipe-separated formats
            row = line.replace("|", "\t").split("\t")
            row = [cell.strip() for cell in row if cell.strip()]  # Clean empty cells
            if row:  # Only add non-empty rows
                coverage_matrix.append(row)
        elif current_table == "validation" and ("\t" in line or "|" in line):
            row = line.replace("|", "\t").split("\t")
            row = [cell.strip() for cell in row if cell.strip()]
            if row:
                validation_table.append(row)
    
    # Create default data if no valid data was parsed
    if not coverage_matrix:
        coverage_matrix = [["Key Parameter", "Covered by Question(s)?", "Question ID(s)", "Gap?", "Comments"],
                         ["No data available", "N/A", "N/A", "N/A", "No data parsed from analysis"]]
    
    if not validation_table:
        validation_table = [["Question ID", "Aligned with C40?", "Key Parameters Missing", "Comments"],
                          ["No data available", "N/A", "N/A", "No data parsed from analysis"]]
    
    # Ensure we have headers
    if len(coverage_matrix) == 1:
        coverage_matrix.append(["No data available", "N/A", "N/A", "N/A", "No data parsed from analysis"])
    if len(validation_table) == 1:
        validation_table.append(["No data available", "N/A", "N/A", "No data parsed from analysis"])
    
    # Create Excel writer object
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        # Create Coverage Matrix sheet
        df_coverage = pd.DataFrame(coverage_matrix[1:], columns=coverage_matrix[0])
        df_coverage.to_excel(writer, sheet_name='Coverage Matrix', index=False)
        
        # Create Validation Table sheet
        df_validation = pd.DataFrame(validation_table[1:], columns=validation_table[0])
        df_validation.to_excel(writer, sheet_name='Validation Table', index=False)

# --- Main Execution ---
if __name__ == "__main__":
    api_key = "AIzaSyDZ1hkeltOGVCVMT6h_lRZGNpyfIgwDOeY"  # Replace with your actual Gemini API key
    input_folder = "D:/Prompt_validation/Prompt_analysis/C40ClimateActionPlanGuide"  # The folder containing your C40 PDF files

    print("🚀 Starting batch C40 CAP analysis...")
    getScores(api_key, input_folder)
    print("✅ All files processed.")