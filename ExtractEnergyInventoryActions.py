import google.generativeai as genai
from pathlib import Path
import time
import os # Added for potentially getting API key from environment
from docx import Document
import re

# --- Function Definition ---
def getScores(api_key: str, folder_path: str) -> None:
    """
    Goes through all PDF files in the specified folder, analyzes each using Gemini
    with a fixed prompt, and saves the responses to text files.

    Args:
        api_key: Your Google Gemini API key.
        folder_path: Path to the folder containing PDF files.
    """
    try:
        genai.configure(api_key=api_key)
        print("Gemini API configured.")
    except Exception as e:
        print(f"Error configuring Gemini API: {e}")
        return

    folder = Path(folder_path.strip())
    if not folder.is_dir():
        print(f"Error: Folder not found at '{folder_path}'")
        return

    # Create analysis_output folder
    output_folder = Path(folder) / "analysis_output"
    output_folder.mkdir(exist_ok=True)
    print(f"Created/accessed output folder: {output_folder}")

    pdf_files = list(folder.glob("*.pdf"))
    if not pdf_files:
        print("No PDF files found in the specified folder.")
        return

       # 7. Define the analysis prompt
    prompt = """
You are a climate action planning expert.

From the provided document, extract only the actions that relate to the emissions inventory. These should include any activity, program, or measure that supports:

GHG emissions data collection or management

GPC protocol adoption

Scope 1, 2, or 3 emissions tracking

Sector- or fuel-type disaggregation

Methodologies or tools for emissions inventory

Validation or verification of emissions data

Any planned updates to the emissions inventory

Output:
List only the actions related to emissions inventory. Do not include actions related to adaptation, finance, or unrelated sectors. Keep each action as a concise bullet point.
"""

    model = genai.GenerativeModel(model_name="gemini-2.5-pro-preview-03-25")

    for pdf_path in pdf_files:
        print(f"\nProcessing '{pdf_path.name}'...")
        pdf_file = None

        try:
            print("Uploading file to Gemini...")
            pdf_file = genai.upload_file(path=pdf_path)
            print(f"Successfully uploaded '{pdf_path.name}' as file ID: {pdf_file.name}")

            print("Waiting for file processing...")
            while pdf_file.state.name == "PROCESSING":
                time.sleep(5)
                pdf_file = genai.get_file(name=pdf_file.name)
                print(f"Current file state: {pdf_file.state.name}")

            if pdf_file.state.name != "ACTIVE":
                print(f"Error: File processing failed or file is not active.")
                print(f"Final state: {pdf_file.state.name}")
                try:
                    genai.delete_file(name=pdf_file.name)
                    print(f"Cleaned up file {pdf_file.name} from server.")
                except Exception as delete_err:
                    print(f"Note: Could not delete file {pdf_file.name} after processing failure: {delete_err}")
                continue

            print("File processed and ready for analysis.")
            print("Sending request to Gemini for analysis...")
            response = model.generate_content([prompt, pdf_file])

            if response.text:
                output_file = output_folder / f"{pdf_path.stem}_analysis.docx"
                doc = Document()
                doc.add_heading(f"Analysis for {pdf_path.name}", level=1)
                bold_pattern = re.compile(r"\*\*(.+?)\*\*")
                for line in response.text.strip().split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    p = doc.add_paragraph()
                    last_idx = 0
                    for match in bold_pattern.finditer(line):
                        if match.start() > last_idx:
                            p.add_run(line[last_idx:match.start()])
                        bold_run = p.add_run(match.group(1))
                        bold_run.bold = True
                        last_idx = match.end()
                    if last_idx < len(line):
                        p.add_run(line[last_idx:])
                doc.save(output_file)
                print(f"Analysis saved to: {output_file.name}")
            else:
                print("[No text content received in the response]")

        except Exception as e:
            print(f"\nAn error occurred during processing or analysis: {e}")

        finally:
            if pdf_file and hasattr(pdf_file, 'name'):
                try:
                    genai.delete_file(name=pdf_file.name)
                    print(f"\nDeleted file {pdf_file.name} from Gemini server.")
                except Exception as cleanup_error:
                    print(f"\nWarning: Failed to delete file {pdf_file.name} from Gemini server: {cleanup_error}")
                    print("You may need to delete it manually via the API or console.")

# --- Example Usage (How to call the function) ---
if __name__ == "__main__":
    my_api_key = "AIzaSyDZ1hkeltOGVCVMT6h_lRZGNpyfIgwDOeY"
    folder_path_input = "D:/All Municipality Reports/Final_Municipality_Reports_For_Analysis/ExtractEnergyEmissions" #"test_folder" #input("Please enter the full path to the folder containing PDF files: ")

    if my_api_key:
        print("\nStarting batch analysis process...")
        getScores(my_api_key, folder_path_input)
        print("\nBatch processing completed.")
    else:
        print("Cannot proceed without a Gemini API Key.")