import google.generativeai as genai
from pathlib import Path
import time
import os
from docx import Document
import re

def getScores(api_key: str, folder_path: str) -> None:
    """
    Processes all PDFs in the folder using Gemini and saves structured DOCX files.
    """
    # 1. Authenticate with Gemini
    try:
        genai.configure(api_key=api_key)
        print("✅ Gemini API configured.")
    except Exception as e:
        print(f"❌ Error configuring Gemini API: {e}")
        return

    # 2. Check input folder
    folder = Path(folder_path.strip())
    if not folder.is_dir():
        print(f"❌ Folder not found: '{folder_path}'")
        return

    # 3. Output folder
    output_folder = folder / "analysis_output"
    output_folder.mkdir(exist_ok=True)
    print(f"📁 Output folder: {output_folder}")

    # 4. Get PDF files
    pdf_files = list(folder.glob("*.pdf"))
    if not pdf_files:
        print("⚠️ No PDF files found.")
        return

    # 5. Gemini prompt
    prompt = """
You are an expert validator analyzing Climate Action Plans (CAPs) and C40 Framework requirements. Your task is to perform a comprehensive content validation between C40 guidance documents and assessment criteria.

OBJECTIVE:
Perform a detailed analysis of C40 Climate Action Planning Framework documents to:
1. Extract all key requirements, guidelines, and best practices
2. Create a comprehensive content validation matrix
3. Identify any gaps in the assessment questions

ANALYSIS STEPS:

1. CONTENT EXTRACTION
- Read the provided C40 PDF thoroughly
- Extract ALL key requirements, guidelines, criteria, and best practices
- Categorize them under relevant themes (e.g., Emissions Inventory, Strategy Development, Implementation, Monitoring)
- Note specific technical requirements, methodologies, or standards mentioned

2. SYSTEMATIC VALIDATION
For each extracted requirement:
- Map it against the 23 assessment questions
- Determine if it's fully, partially, or not covered
- Note specific question numbers that address each requirement
- Identify any missing aspects or gaps

3. OUTPUT GENERATION

Generate three detailed matrices:

Matrix 1: C40 Requirements Coverage
| C40 Requirement | Source Section | Requirement Type | Covered by Question(s)? | Question ID(s) | Coverage Level | Notes |
Example:
| GPC Protocol Usage | Emissions Inventory | Mandatory | Yes | 1.1 | Full | Explicitly covered in inventory question |

Matrix 2: Assessment Questions Validation
| Question ID | Question Text | C40 Requirements Covered | Missing Requirements | Suggested Improvements |
Example:
| 1.1 | GPC adoption question | GPC protocol requirement | None | Consider adding version specification |

Matrix 3: Gap Analysis
| Gap Area | Missing C40 Requirements | Importance Level | Suggested New Questions |
Example:
| Data Quality | Independent verification requirement | High | Add question about third-party verification |

ASSESSMENT QUESTIONS FOR REFERENCE:
1. Emissions Inventory 
1.1 Has the city explicitly adopted the GPC (BASIC or BASIC+) to structure its emissions inventory? 
1.2 Does the city's emissions inventory explicitly include Scope 1 stationary energy emissions (on-site fuel combustion)? 
1.3 Does the city's emissions inventory explicitly include Scope 2 stationary energy emissions (grid-supplied electricity, heat, or steam)? 
1.4 Does the city's emissions inventory explicitly state whether Scope 3 stationary energy emissions (e.g., upstream energy losses, imported fuels) are included or not? 
1.5 Does the emissions inventory explicitly disaggregate stationary energy emissions by sub-sector (e.g., residential, commercial, industrial)?
1.6 Does the emissions inventory explicitly disaggregate stationary energy emissions by fuel type (e.g., electricity, natural gas, coal, renewables)? 
2. Strategy Identification 
2.1 Does the city explicitly mention using emissions modeling or scenario planning for stationary energy emissions reduction strategies? 
2.2	Does the city explicitly identify technological solutions (e.g., renewable energy, retrofitting, heat pumps)? 

2.3	Does the city explicitly identify non-technological solutions (e.g., policy changes, behavior change programs)? 

2.4	Has the city explicitly set GHG reduction targets specifically for stationary energy sub-sectors (e.g., residential buildings, commercial buildings, industries, or energy utilities)? 

2.5	Has the city explicitly addressed residual stationary energy emissions (e.g., offsets, negative emission strategies)? 

2.6	Has the city explicitly assessed how climate risks (e.g., heatwaves, flooding) impact the stationary energy sector (demand or reliability)? 

2.7	Has the city explicitly identified adaptation strategies (e.g., resilient infrastructure, grid reliability enhancements) addressing climate risks specific to the stationary energy sector? 

3. Action Prioritization & Detailing 
3.1	Does the city explicitly prioritize stationary energy actions based on quantified or qualitatively stated GHG emissions reduction potential? 

3.2	Does the city explicitly document feasibility assessments (technical, financial, institutional, or political) for prioritized stationary energy actions? 

3.3	Does the city explicitly address equity by considering vulnerable communities (e.g., low-income housing) when prioritizing stationary energy actions?

3.4	For each prioritized stationary energy action, does the city explicitly define the expected GHG or climate adaptation impact? 

3.5	Does the city explicitly identify responsible departments or governance structures for implementing prioritized stationary energy actions?

4	Monitoring, Evaluation & Reporting (MER)

4.1	Does the city explicitly establish Key Performance Indicators (KPIs) to monitor stationary energy actions (e.g., energy savings, buildings retrofitted, emissions reduced)? 

4.2	Does the city explicitly describe regular evaluations (annual, biennial, mid-term) of stationary energy actions?

4.3	Does the city explicitly commit to publicly reporting progress on stationary energy actions (e.g., annual reports, dashboards)? 

4.4	Does the city explicitly state that monitoring data is used to revise or update the CAP or stationary energy strategies? 

4.5	Does the city explicitly describe data quality assurance procedures (validation or verification of stationary energy emission data)? 

SPECIFIC VALIDATION CRITERIA:
1. Completeness: Does the question capture all aspects of the C40 requirement?
2. Specificity: Is the question specific enough to assess compliance?
3. Measurability: Can the response be objectively evaluated?
4. Relevance: Does the question align with C40's intent?

IMPORTANT CONSIDERATIONS:
- Focus on both explicit and implicit requirements in C40 documents
- Consider technical, procedural, and governance aspects
- Note any regional or city-size specific requirements
- Identify both mandatory and recommended practices
- Consider alignment with international standards referenced by C40

Please provide detailed, actionable feedback that can be used to improve the assessment framework.

    """

    # 6. Load Gemini model
    model = genai.GenerativeModel("gemini-2.5-pro-preview-03-25")

    # 7. Process each PDF
    for pdf_path in pdf_files:
        print(f"\n📄 Processing: {pdf_path.name}")
        pdf_file = None

        try:
            pdf_file = genai.upload_file(path=pdf_path)
            print(f"✅ Uploaded: {pdf_file.name}")

            while pdf_file.state.name == "PROCESSING":
                time.sleep(5)
                pdf_file = genai.get_file(name=pdf_file.name)
                print(f"⏳ Waiting... Current state: {pdf_file.state.name}")

            if pdf_file.state.name != "ACTIVE":
                print(f"❌ File not active. Final state: {pdf_file.state.name}")
                continue

            print("🧠 Sending prompt to Gemini...")
            response = model.generate_content([prompt, pdf_file])

            if response.text:
                output_file = output_folder / f"{pdf_path.stem}_analysis.docx"
                doc = Document()
                doc.add_heading(f"Analysis for {pdf_path.name}", level=1)

                bold_pattern = re.compile(r"\*\*(.+?)\*\*")
                for line in response.text.strip().split("\n"):
                    if not line.strip():
                        continue
                    p = doc.add_paragraph()
                    last_idx = 0
                    for match in bold_pattern.finditer(line):
                        if match.start() > last_idx:
                            p.add_run(line[last_idx:match.start()])
                        bold_run = p.add_run(match.group(1))
                        bold_run.bold = True
                        last_idx = match.end()
                    if last_idx < len(line):
                        p.add_run(line[last_idx:])
                doc.save(output_file)
                print(f"✅ Saved: {output_file.name}")
            else:
                print("⚠️ No response text from Gemini.")

        except Exception as e:
            print(f"❌ Error: {e}")

        finally:
            if pdf_file and hasattr(pdf_file, 'name'):
                try:
                    genai.delete_file(name=pdf_file.name)
                    print(f"🧹 Deleted temp file: {pdf_file.name}")
                except Exception as cleanup_error:
                    print(f"⚠️ Failed to delete: {pdf_file.name} - {cleanup_error}")

# ---- Run Script ----
if __name__ == "__main__":
    api_key = "AIzaSyDZ1hkeltOGVCVMT6h_lRZGNpyfIgwDOeY"  # Replace this with a secure method (e.g., environment variable)
    folder_path = "D:/Prompt_validation/Prompt_analysis/C40ClimateActionPlanGuide"

    if api_key and folder_path:
        print("🚀 Starting Gemini batch analysis...")
        getScores(api_key, folder_path)
        print("✅ All PDFs processed.")
    else:
        print("❌ API key or folder path missing.")
