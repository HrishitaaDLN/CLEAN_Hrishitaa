import google.generativeai as genai
from pathlib import Path
import time
import os
import csv
from datetime import datetime
from docx import Document
from docx.shared import Pt
import re


def analyze_reports_in_folder(folder_path: str, api_key: str):
    """
    Analyzes all PDF files in a given folder to determine their type and relevance
    to sustainability/emissions using Gemini, and summarizes key metadata.

    Args:
        folder_path: Path to folder containing PDF files.
        api_key: Gemini API key.

    Returns:
        List of summary dictionaries for each processed file.
    """
    try:
        genai.configure(api_key=api_key)
        print("Gemini API configured.")
    except Exception as e:
        print(f"Failed to configure Gemini API: {e}")
        return []

    folder = Path(folder_path)
    if not folder.is_dir():
        print(f"Invalid folder path: {folder_path}")
        return []

    summaries = []
    pdf_files = list(folder.glob("*.pdf"))
    if not pdf_files:
        print("No PDF files found in the folder.")
        return []

    prompt = """
        You are given a city or village level report in PDF format.
        Determine if this report is relevant to climate change, emissions, or sustainability per C40 context.

        Provide a concise and concrete summary with the following structure:
        1. **Report Name**: Extract title and file name based on the report and its discussion regarding sustainability context.
        2. **Community/Village/City Name**: Identify the location the report is about.
        3. **Report Type**: E.g., GHG Inventory, Climate Action Plan, Adaptation Assessment, or unrelated.
        4. **Sustainability Content Summary**: 2-3 line overview of what this report contains regarding sustainability.
        5. **Published Date**: If found in the document.

        Focus strictly on clarity, accuracy, and completeness of above items.
    """

    model = genai.GenerativeModel(model_name="gemini-2.5-pro-preview-03-25")

    for pdf_path in pdf_files:
        print(f"\nProcessing: {pdf_path.name}")
        try:
            pdf_file = genai.upload_file(path=pdf_path)
            print(f"Uploaded: {pdf_path.name} as ID {pdf_file.name}")

            while pdf_file.state.name == "PROCESSING":
                time.sleep(5)
                pdf_file = genai.get_file(name=pdf_file.name)

            if pdf_file.state.name != "ACTIVE":
                print(f"Skipping: {pdf_path.name} (Not active after upload)")
                continue

            response = model.generate_content([prompt, pdf_file])
            if response.text:
                
                summaries.append({
                    "file": pdf_path.name,
                    "summary": response.text.strip()
                })
            else:
                print(f"No response for: {pdf_path.name}")

        except Exception as e:
            print(f"Error processing {pdf_path.name}: {e}")
        finally:
            try:
                genai.delete_file(name=pdf_file.name)
                print(f"Deleted file: {pdf_file.name}")
            except:
                print(f"Cleanup failed for file: {pdf_path.name}")

    return summaries

def save_summaries_to_csv(summaries: list, output_path: str = None):
    """
    Saves the report summaries to a CSV file.
    
    Args:
        summaries: List of summary dictionaries
        output_path: Optional path for the CSV file. If not provided, creates a timestamped file in current directory.
    """
    if not summaries:
        print("No summaries to save.")
        return
        
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"report_summaries_{timestamp}.csv"
    
    # Convert output_path to Path object
    output_path = Path(output_path)
    
    try:
        # Create parent directory if it doesn't exist
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['file', 'summary']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            
            writer.writeheader()
            for summary in summaries:
                writer.writerow(summary)
                
        print(f"Summaries saved to: {output_path}")
    except Exception as e:
        print(f"Error saving CSV file: {e}")

def save_summaries_to_txt(summaries: list, output_path: str = None):
    """
    Saves the report summaries to a text file with formatted output.
    
    Args:
        summaries: List of summary dictionaries
        output_path: Optional path for the text file. If not provided, creates a timestamped file in current directory.
    """
    if not summaries:
        print("No summaries to save.")
        return
        
    if output_path is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"report_summaries_{timestamp}.txt"
    
    # Convert output_path to Path object
    output_path = Path(output_path)
    
    try:
        # Create parent directory if it doesn't exist
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as txtfile:
            for summary in summaries:
                txtfile.write(f"\n{'='*80}\n")
                txtfile.write(f"File: {summary['file']}\n")
                txtfile.write(f"{'='*80}\n\n")
                txtfile.write(summary['summary'])
                txtfile.write("\n\n")
                
        print(f"Summaries saved to text file: {output_path}")
    except Exception as e:
        print(f"Error saving text file: {e}")

def save_summaries_to_docx(summaries: list, output_folder: str = None):
    """
    Saves all report summaries into a single .docx file with improved formatting.
    Args:
        summaries: List of summary dictionaries
        output_folder: Optional output folder path. Defaults to current directory.
    """
    if not summaries:
        print("No summaries to save.")
        return

    if output_folder is None:
        output_folder = "docx_outputs"
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("All Report Summaries", level=0)

    bold_pattern = re.compile(r"\*\*(.+?)\*\*")

    for summary in summaries:
        doc.add_heading(f"Summary: {summary['file']}", level=1)
        for line in summary['summary'].splitlines():
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            last_idx = 0
            for match in bold_pattern.finditer(line):
                # Add text before bold
                if match.start() > last_idx:
                    p.add_run(line[last_idx:match.start()])
                # Add bold text
                bold_run = p.add_run(match.group(1))
                bold_run.bold = True
                bold_run.font.size = Pt(11)
                last_idx = match.end()
            # Add any remaining text after the last bold
            if last_idx < len(line):
                p.add_run(line[last_idx:])
        doc.add_paragraph("")  # Add space between summaries

    filename = output_path / "all_summaries.docx"
    doc.save(filename)
    print(f"Saved all summaries to DOCX: {filename}")


# Example usage:
if __name__ == "__main__":
    API_KEY = "AIzaSyDZ1hkeltOGVCVMT6h_lRZGNpyfIgwDOeY"
    FOLDER_PATH = "C:/Users/Admin/Downloads/All Municipality Reports-selected"

    results = analyze_reports_in_folder(FOLDER_PATH, API_KEY)
    
    # Print summaries to console
    for item in results:
        print(f"\n--- Summary for {item['file']} ---\n{item['summary']}")
    
    # Save summaries to both CSV and text formats
    # Create analysis_output folder
    output_folder = Path(FOLDER_PATH) / "analysis_output"
    output_folder.mkdir(exist_ok=True)
    print(f"Created/accessed output folder: {output_folder}")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    save_summaries_to_csv(results, output_folder / f"report_summaries_{timestamp}.csv")
    save_summaries_to_txt(results, output_folder / f"report_summaries_{timestamp}.txt")
    save_summaries_to_docx(results, output_folder)